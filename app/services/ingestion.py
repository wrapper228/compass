from __future__ import annotations

import csv as _csv
import hashlib
import io
import json
import re
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, List, Optional

import structlog
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy import func, select
from sqlalchemy.orm import Session, selectinload

from app.core.config import get_settings
from app.db import models
from app.db.session import SessionLocal
from app.services.retrieval.chunker import Chunk, chunk_document, deduplicate_chunks
from app.services.retrieval.index_manager import ChunkIndexInput, HybridIndexManager

logger = structlog.get_logger(__name__)


@dataclass
class SourceDocument:
    relative_path: str
    text: str
    sha256: str
    file_size: int
    name: str
    folder: str


def process_zip(
    job_id: str,
    zip_bytes: bytes,
    workspace: Path,
    dataset_slug: str,
    dataset_title: Optional[str] = None,
    dataset_description: Optional[str] = None,
) -> dict:
    settings = get_settings()
    slug = slugify_dataset(dataset_slug)
    title = dataset_title.strip() if dataset_title else slug.replace("-", " ").title()

    base_dir = workspace / slug / job_id
    base_dir.mkdir(parents=True, exist_ok=True)
    raw_dir = base_dir / "raw"
    raw_dir.mkdir(exist_ok=True)
    result_path = base_dir / "ingest_result.json"

    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as archive:
        archive.extractall(raw_dir)

    documents = list(_collect_documents(raw_dir))
    if not documents:
        result = {
            "job_id": job_id,
            "dataset": slug,
            "documents": 0,
            "chunks": 0,
            "indexed": 0,
        }
        result_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
        logger.info("ingest.no_documents", job_id=job_id, dataset=slug)
        return result

    manager = HybridIndexManager(settings)
    session = SessionLocal()
    indexed_inputs: List[ChunkIndexInput] = []
    chunk_total = 0
    document_total = 0

    logger.info(
        "ingest.started",
        job_id=job_id,
        dataset=slug,
        documents=len(documents),
    )

    try:
        dataset = _ensure_dataset(session, slug, title, dataset_description)

        for doc in documents:
            chunks = _chunk_document(doc, settings.chunk_size_tokens, settings.chunk_overlap_tokens)
            if not chunks:
                continue

            for idx, chunk in enumerate(chunks):
                chunk.ordinal = idx

            removed_chunk_ids = _delete_existing_document(session, dataset.id, doc.relative_path)
            if removed_chunk_ids and settings.retrieval_enabled:
                manager.remove_chunks(removed_chunk_ids, source=f"dataset:{dataset.slug}", session=session)

            doc_row = models.KnowledgeDocument(
                dataset_id=dataset.id,
                path=doc.relative_path,
                name=doc.name,
                folder=doc.folder,
                sha256=doc.sha256,
                file_size=doc.file_size,
                token_count=sum(chunk.token_count for chunk in chunks),
                chunk_count=len(chunks),
            )
            session.add(doc_row)
            session.flush()

            summary_text, tail_text = _build_summaries(chunks)
            session.add(
                models.KnowledgeDocumentSummary(
                    document_id=doc_row.id,
                    summary_text=summary_text,
                    tail_text=tail_text,
                )
            )

            chunk_rows: List[tuple[models.KnowledgeChunk, Chunk]] = []
            for chunk in chunks:
                chunk_row = models.KnowledgeChunk(
                    document_id=doc_row.id,
                    ordinal=chunk.ordinal,
                    text=chunk.text,
                    start_line=chunk.start_line,
                    end_line=chunk.end_line,
                    char_start=chunk.char_start,
                    char_end=chunk.char_end,
                    token_count=chunk.token_count,
                    sha256=chunk.sha256,
                )
                session.add(chunk_row)
                chunk_rows.append((chunk_row, chunk))

            session.flush()

            for row, chunk in chunk_rows:
                indexed_inputs.append(
                    ChunkIndexInput(
                        chunk_id=row.id,
                        text=chunk.text,
                        document_path=doc.relative_path,
                        document_sha=doc.sha256,
                        start_line=chunk.start_line,
                        end_line=chunk.end_line,
                        dataset_slug=dataset.slug,
                        folder=doc.folder,
                    )
                )

            chunk_total += len(chunk_rows)
            document_total += 1

        session.flush()

        indexed = 0
        if indexed_inputs and settings.retrieval_enabled:
            indexed = manager.index_chunks(indexed_inputs, source=f"dataset:{dataset.slug}", session=session)

        # refresh dataset totals
        dataset.total_documents = session.query(func.count(models.KnowledgeDocument.id)).filter(
            models.KnowledgeDocument.dataset_id == dataset.id
        ).scalar()
        dataset.total_chunks = session.query(func.count(models.KnowledgeChunk.id)).join(
            models.KnowledgeDocument
        ).filter(models.KnowledgeDocument.dataset_id == dataset.id).scalar()
        dataset.total_tokens = (
            session.query(func.coalesce(func.sum(models.KnowledgeChunk.token_count), 0))
            .join(models.KnowledgeDocument)
            .filter(models.KnowledgeDocument.dataset_id == dataset.id)
            .scalar()
        )
        dataset.last_ingested_at = datetime.utcnow()

        session.commit()
        logger.info(
            "ingest.completed",
            job_id=job_id,
            dataset=dataset.slug,
            documents=document_total,
            chunks=chunk_total,
            indexed=indexed,
        )
    except Exception:
        session.rollback()
        if settings.retrieval_enabled:
            try:
                manager.rebuild(source="rollback")
            except Exception:
                logger.exception("ingest.rebuild_failed", dataset=slug)
        logger.exception("ingest.failed", dataset=slug)
        raise
    finally:
        session.close()

    result = {
        "job_id": job_id,
        "dataset": slug,
        "documents": document_total,
        "chunks": chunk_total,
        "indexed": indexed,
    }
    result_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")

    try:
        from app.services.retrieval import refresh_indices

        refresh_indices()
    except Exception:
        logger.exception("ingest.refresh_failed", dataset=slug)

    return result


def _ensure_dataset(
    session: Session,
    slug: str,
    title: str,
    description: Optional[str],
) -> models.KnowledgeDataset:
    dataset = session.execute(
        select(models.KnowledgeDataset).where(models.KnowledgeDataset.slug == slug)
    ).scalar_one_or_none()
    if dataset:
        if description:
            dataset.description = description
        dataset.title = title
        return dataset

    dataset = models.KnowledgeDataset(
        slug=slug,
        title=title,
        description=description,
        created_at=datetime.utcnow(),
    )
    session.add(dataset)
    session.flush()
    return dataset


def _collect_documents(raw_dir: Path) -> Iterable[SourceDocument]:
    for path in sorted(raw_dir.rglob("*")):
        if not path.is_file():
            continue
        rel_path = path.relative_to(raw_dir).as_posix()
        text = _extract_text(path)
        if not text:
            continue
        sha = hashlib.sha256(text.encode("utf-8")).hexdigest()
        folder = path.parent.relative_to(raw_dir).as_posix() if path.parent != raw_dir else ""
        yield SourceDocument(
            relative_path=rel_path,
            text=text,
            sha256=sha,
            file_size=path.stat().st_size,
            name=path.name,
            folder=folder,
        )


def _extract_text(path: Path) -> Optional[str]:
    ext = path.suffix.lower()
    try:
        if ext in {".txt", ".md"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        if ext == ".json":
            data = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
            return json.dumps(data, ensure_ascii=False, indent=2)
        if ext == ".csv":
            rows: List[str] = []
            with path.open("r", encoding="utf-8", errors="ignore") as fh:
                reader = _csv.reader(fh)
                for row in reader:
                    rows.append(", ".join(row))
            return "\n".join(rows)
        if ext in {".html", ".htm"}:
            html = path.read_text(encoding="utf-8", errors="ignore")
            soup = BeautifulSoup(html, "lxml")
            return soup.get_text("\n", strip=True)
        if ext == ".docx":
            doc = DocxDocument(path)
            return "\n".join(para.text for para in doc.paragraphs)
        if ext == ".pdf":
            reader = PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n".join(pages)
    except Exception:
        return None
    return None


def _chunk_document(doc: SourceDocument, chunk_size: int, overlap: int) -> List[Chunk]:
    chunks = chunk_document(doc.text, chunk_size=chunk_size, overlap=overlap)
    return deduplicate_chunks(chunks)


def _delete_existing_document(session: Session, dataset_id: int, relative_path: str) -> List[int]:
    stmt = (
        select(models.KnowledgeDocument)
        .options(selectinload(models.KnowledgeDocument.chunks))
        .where(
            models.KnowledgeDocument.dataset_id == dataset_id,
            models.KnowledgeDocument.path == relative_path,
        )
    )
    doc = session.execute(stmt).scalar_one_or_none()
    if doc is None:
        return []
    chunk_ids = [chunk.id for chunk in doc.chunks]
    session.delete(doc)
    session.flush()
    return chunk_ids


def _build_summaries(chunks: List[Chunk]) -> tuple[str, Optional[str]]:
    if not chunks:
        return "", None
    head = chunks[0].text.strip()
    tail = " ".join(chunk.text.strip() for chunk in chunks[-2:]) if len(chunks) > 1 else chunks[-1].text.strip()
    head_trimmed = head[:600].strip()
    tail_trimmed = tail[:800].strip()
    return head_trimmed, tail_trimmed


_slug_pattern = re.compile(r"[^a-z0-9\-]+")


def slugify_dataset(value: str) -> str:
    normalized = value.strip().lower().replace(" ", "-")
    normalized = _slug_pattern.sub("-", normalized)
    normalized = normalized.strip("-")
    return normalized or "dataset"

